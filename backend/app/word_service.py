"""Word document generation service for commercial offers."""

from datetime import datetime, timezone
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# Constants
OFFER_VALIDITY_DAYS = 30


def remove_diacritics(text):
    """
    Remove Romanian diacritics from text.
    Converts: ă→a, â→a, î→i, ș→s, ț→t (both uppercase and lowercase)
    """
    if not text:
        return text
    
    if not isinstance(text, str):
        text = str(text)
    
    diacritics_map = {
        'ă': 'a', 'Ă': 'A',
        'â': 'a', 'Â': 'A',
        'î': 'i', 'Î': 'I',
        'ș': 's', 'Ș': 'S',
        'ț': 't', 'Ț': 'T'
    }
    
    for diacritic, replacement in diacritics_map.items():
        text = text.replace(diacritic, replacement)
    
    return text


def set_cell_background(cell, color):
    """Set background color for a table cell."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)


def add_heading_with_style(doc, text, level=1):
    """Add a heading with custom styling."""
    heading = doc.add_heading(remove_diacritics(text), level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(30, 64, 175)  # Blue color
    return heading


def add_two_column_table(doc, data):
    """Add a two-column table with label and value."""
    table = doc.add_table(rows=len(data), cols=2)
    table.style = 'Light Grid Accent 1'
    
    for i, (label, value) in enumerate(data):
        row = table.rows[i]
        row.cells[0].text = remove_diacritics(str(label))
        row.cells[1].text = remove_diacritics(str(value))
        
        # Make label bold
        for paragraph in row.cells[0].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    return table


def add_materials_table(doc, materials_list):
    """Add materials table with headers and data."""
    # Add table with header row
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    header_cells = table.rows[0].cells
    headers = ['Nr.', 'Material', 'SKU', 'Cantitate', 'Pret Unitar cu Adaos (RON)', 'Total (RON)']
    
    for i, header in enumerate(headers):
        cell = header_cells[i]
        cell.text = remove_diacritics(header)
        set_cell_background(cell, '1e40af')
        
        # Make header text white and bold
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
    
    # Data rows
    total_cost = 0
    for idx, material in enumerate(materials_list, 1):
        material_total = material.get('quantity_planned', 0) * material.get('unit_price', 0)
        total_cost += material_total
        
        row = table.add_row()
        cells = row.cells
        cells[0].text = str(idx)
        cells[1].text = remove_diacritics(material.get('material_name', 'N/A'))
        cells[2].text = remove_diacritics(material.get('material_sku', 'N/A'))
        cells[3].text = f"{material.get('quantity_planned', 0):.2f}"
        cells[4].text = f"{material.get('unit_price', 0):.2f}"
        cells[5].text = f"{material_total:.2f}"
        
        # Align numeric columns to the right
        for i in [0, 3, 4, 5]:
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Add subtotal row
    subtotal_row = table.add_row()
    subtotal_cells = subtotal_row.cells
    
    # Merge first 5 cells
    merged_cell = subtotal_cells[0].merge(subtotal_cells[4])
    merged_cell.text = remove_diacritics('SUBTOTAL MATERIALE (cu adaos):')
    merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    subtotal_cells[5].text = f"{total_cost:.2f}"
    subtotal_cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Make subtotal row bold and add background
    for cell in [merged_cell, subtotal_cells[5]]:
        set_cell_background(cell, 'dbeafe')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    return table, total_cost


def add_additional_costs_table(doc, labor_cost, transport_cost, other_costs):
    """Add additional costs table."""
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = remove_diacritics('Tip Cost')
    header_cells[1].text = remove_diacritics('Valoare (RON)')
    
    for cell in header_cells:
        set_cell_background(cell, '1e40af')
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
    
    # Cost rows
    costs = [
        ('Manopera', labor_cost),
        ('Transport', transport_cost),
        ('Alte costuri', other_costs)
    ]
    
    for label, value in costs:
        row = table.add_row()
        cells = row.cells
        cells[0].text = remove_diacritics(label)
        cells[1].text = f"{value:.2f}"
        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Add subtotal row
    subtotal = labor_cost + transport_cost + other_costs
    subtotal_row = table.add_row()
    cells = subtotal_row.cells
    cells[0].text = remove_diacritics('SUBTOTAL COSTURI ADITIONALE:')
    cells[1].text = f"{subtotal:.2f}"
    
    for cell in cells:
        set_cell_background(cell, 'dbeafe')
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    return table, subtotal


def add_grand_total_table(doc, total_amount):
    """Add grand total table."""
    table = doc.add_table(rows=1, cols=2)
    
    cells = table.rows[0].cells
    cells[0].text = remove_diacritics('TOTAL GENERAL:')
    cells[1].text = f"{total_amount:.2f} RON"
    
    for cell in cells:
        set_cell_background(cell, '1e40af')
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(255, 255, 255)
    
    return table


def extract_additional_costs(project_data):
    """
    Extract additional costs from project data.
    
    Args:
        project_data: Dictionary containing project information
    
    Returns:
        tuple: (labor_cost, transport_cost, other_costs)
    """
    labor_cost = project_data.get('labor_cost_estimated', 0) or 0
    transport_cost = project_data.get('transport_cost_estimated', 0) or 0
    other_costs = project_data.get('other_costs_estimated', 0) or 0
    return labor_cost, transport_cost, other_costs


def generate_commercial_offer_word(project_data, materials_list=None):
    """
    Generate a commercial offer Word document for a project.
    
    Args:
        project_data: Dictionary containing project information
        materials_list: List of materials with quantities and prices
    
    Returns:
        BytesIO: Word document as bytes
    """
    doc = Document()
    
    # Set document margins (narrower margins)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
    
    # Title
    title = doc.add_heading(remove_diacritics('OFERTA COMERCIALA'), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(30, 64, 175)
    
    doc.add_paragraph()  # Spacer
    
    # Offer Details Section
    add_heading_with_style(doc, 'Detalii Oferta', level=2)
    
    now = datetime.now(timezone.utc)
    offer_date = now.strftime("%d.%m.%Y")
    offer_details = [
        ('Data ofertei:', offer_date),
        ('Nr. oferta:', f"OF-{project_data.get('id', 'N/A')}-{now.strftime('%Y%m%d')}"),
        ('Valabilitate:', f'{OFFER_VALIDITY_DAYS} zile')
    ]
    
    add_two_column_table(doc, offer_details)
    doc.add_paragraph()  # Spacer
    
    # Client Information
    add_heading_with_style(doc, 'Date Client', level=2)
    
    client_info = [
        ('Nume client:', project_data.get('client_name', 'N/A')),
        ('Contact:', project_data.get('client_contact', 'N/A')),
        ('Locatie:', project_data.get('location', 'N/A'))
    ]
    
    add_two_column_table(doc, client_info)
    doc.add_paragraph()  # Spacer
    
    # Project Information
    add_heading_with_style(doc, 'Detalii Proiect', level=2)
    
    project_info = [
        ('Nume proiect:', project_data.get('name', 'N/A')),
        ('Capacitate sistem:', f"{project_data.get('capacity_kw', 0)} kW" if project_data.get('capacity_kw') else 'N/A'),
        ('Status:', project_data.get('status', 'N/A').replace('_', ' ').title())
    ]
    
    if project_data.get('start_date'):
        project_info.append(('Data start estimata:', str(project_data.get('start_date'))))
    
    add_two_column_table(doc, project_info)
    doc.add_paragraph()  # Spacer
    
    # Materials List (if provided)
    if materials_list and len(materials_list) > 0:
        add_heading_with_style(doc, 'Materiale si Costuri', level=2)
        
        materials_table, total_cost = add_materials_table(doc, materials_list)
        doc.add_paragraph()  # Spacer
        
        # Additional Costs Section
        labor_cost, transport_cost, other_costs = extract_additional_costs(project_data)
        
        add_heading_with_style(doc, 'Costuri Aditionale', level=2)
        
        additional_costs_table, subtotal_additional = add_additional_costs_table(
            doc, labor_cost, transport_cost, other_costs
        )
        doc.add_paragraph()  # Spacer
        
        # Grand Total
        grand_total = total_cost + subtotal_additional
        add_grand_total_table(doc, grand_total)
        doc.add_paragraph()  # Spacer
    
    # Pricing section (if no materials provided)
    if not materials_list or len(materials_list) == 0:
        labor_cost, transport_cost, other_costs = extract_additional_costs(project_data)
        
        # Show additional costs if any exist
        if labor_cost > 0 or transport_cost > 0 or other_costs > 0:
            add_heading_with_style(doc, 'Costuri Aditionale', level=2)
            add_additional_costs_table(doc, labor_cost, transport_cost, other_costs)
            doc.add_paragraph()  # Spacer
        
        # Show estimated cost if provided
        if project_data.get('estimated_cost'):
            add_heading_with_style(doc, 'Estimare Cost', level=2)
            cost_data = [
                ('Cost estimat total:', f"{project_data.get('estimated_cost', 0):.2f} RON")
            ]
            add_two_column_table(doc, cost_data)
            doc.add_paragraph()  # Spacer
    
    # Notes
    if project_data.get('notes'):
        add_heading_with_style(doc, 'Note', level=2)
        doc.add_paragraph(remove_diacritics(project_data.get('notes')))
        doc.add_paragraph()  # Spacer
    
    # Footer with Terms
    doc.add_page_break()
    add_heading_with_style(doc, 'Termeni si Conditii', level=2)
    
    terms = [
        "• Preturile sunt exprimate in RON si nu includ TVA.",
        f"• Oferta este valabila {OFFER_VALIDITY_DAYS} de zile de la data emiterii.",
        "• Timpul de livrare va fi confirmat la plasarea comenzii.",
        "• Montajul si punerea in functiune sunt incluse in pret.",
        "• Garantie conform specificatiilor producatorilor."
    ]
    
    for term in terms:
        doc.add_paragraph(remove_diacritics(term))
    
    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer.getvalue()
