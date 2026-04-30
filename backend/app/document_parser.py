"""Document parser service for extracting invoice materials from PDF, DOC, and TXT files."""

import re
import logging
from typing import List, Dict, Optional

# Import PDF and DOC parsing libraries
try:
    from PyPDF2 import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)
COMPANY_HINTS = ("srl", "s.r.l", "sa", "s.a", "srl.", "s.a.", "srl,", "s.a,")
TOTAL_KEYWORDS = (
    "total de plata",
    "total plata",
    "de plata",
    "total general",
    "grand total",
    "total",
)
EMPTY_PARSE_RESULT = {
    'invoice_number': '',
    'invoice_date': '',
    'supplier_name': '',
    'currency': 'RON',
    'total_amount': 0.0,
    'items': []
}


def parse_decimal(value: str) -> Optional[float]:
    """Parse decimal values from mixed locale formats."""
    if value is None:
        return None
    cleaned = value.strip().replace(" ", "")
    cleaned = re.sub(r"[^\d,.\-]", "", cleaned)
    if not cleaned:
        return None

    # If both separators exist, assume the right-most separator is decimal.
    if "," in cleaned and "." in cleaned:
        if cleaned.rfind(",") > cleaned.rfind("."):
            cleaned = cleaned.replace(".", "").replace(",", ".")
        else:
            cleaned = cleaned.replace(",", "")
    elif "," in cleaned:
        # Romanian style decimal separator.
        cleaned = cleaned.replace(".", "").replace(",", ".")

    try:
        return float(cleaned)
    except ValueError:
        return None


def extract_amount_candidates(text: str) -> List[float]:
    """Extract numeric amounts with decimal part."""
    pattern = r"(?<!\d)(?:\d{1,3}(?:[.\s,]\d{3})+|\d+)(?:[.,]\d{2})(?!\d)"
    values: List[float] = []
    for match in re.findall(pattern, text):
        value = parse_decimal(match)
        if value is not None:
            values.append(value)
    return values


def is_likely_supplier_line(line: str) -> bool:
    line_lower = line.lower().strip()
    if len(line_lower) < 3:
        return False
    if any(hint in line_lower for hint in COMPANY_HINTS):
        return True
    if line_lower.startswith("furnizor") or line_lower.startswith("supplier"):
        return True
    return False


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text content from PDF file."""
    if not PDF_AVAILABLE:
        raise Exception("PyPDF2 library not available. Install with: pip install PyPDF2")
    
    try:
        text = ""
        with open(file_path, 'rb') as file:
            reader = PdfReader(file)
            for page in reader.pages:
                text += page.extract_text() + "\n"
        return text
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {e}")
        raise Exception(f"Failed to extract text from PDF: {str(e)}")


def extract_text_from_docx(file_path: str) -> str:
    """Extract text content from DOCX file."""
    if not DOCX_AVAILABLE:
        raise Exception("python-docx library not available. Install with: pip install python-docx")
    
    try:
        doc = Document(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                text += "\n" + "\t".join([cell.text for cell in row.cells])
        
        return text
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {e}")
        raise Exception(f"Failed to extract text from DOCX: {str(e)}")


def extract_text_from_doc(file_path: str) -> str:
    """Extract text content from DOC file (legacy format).
    
    Note: python-docx doesn't support .doc files, only .docx.
    For .doc files, we return an empty string and let the user manually process.
    """
    logger.warning("Legacy .doc format not supported. Please convert to .docx or use XML format.")
    return ""


def extract_text_from_txt(file_path: str) -> str:
    """Extract text content from TXT file."""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    except UnicodeDecodeError:
        # Try with different encoding
        try:
            with open(file_path, 'r', encoding='latin-1') as file:
                return file.read()
        except Exception as e:
            logger.error(f"Error reading TXT file: {e}")
            raise Exception(f"Failed to read TXT file: {str(e)}")


def parse_invoice_materials(text: str) -> Dict:
    """Parse invoice materials from extracted text.
    
    This function attempts to extract:
    - Invoice number
    - Supplier name
    - Invoice date
    - Total amount
    - Line items with description, quantity, unit price, and total
    
    The parsing uses heuristics and pattern matching since the text format
    can vary significantly between different invoice formats.
    """
    result = {
        'invoice_number': '',
        'invoice_date': '',
        'supplier_name': '',
        'currency': 'RON',
        'total_amount': 0.0,
        'items': []
    }
    
    lines = [line.strip() for line in text.split('\n')]
    
    # Pattern matching for common invoice fields
    for i, line in enumerate(lines):
        line_lower = line.lower()
        
        # Extract invoice number
        if not result['invoice_number']:
            invoice_patterns = [
                r'(?:factura\s*(?:fiscala)?|invoice)\s*(?:nr|no|number)?\.?[:\s#-]*([A-Z0-9][A-Z0-9\-/]{2,})',
                r'(?:nr\.?\s*factura)[:\s#-]*([A-Z0-9][A-Z0-9\-/]{2,})',
            ]
            for pattern in invoice_patterns:
                match = re.search(pattern, line, re.IGNORECASE)
                if match:
                    result['invoice_number'] = match.group(1).strip()
                    break
        
        # Extract invoice date
        if not result['invoice_date']:
            date_patterns = [
                r'(?:date|data)[:\s]*(\d{1,2}[\./-]\d{1,2}[\./-]\d{2,4})',
                r'(\d{4}-\d{2}-\d{2})',
                r'(\d{1,2}[\./-]\d{1,2}[\./-]\d{2,4})',
            ]
            for pattern in date_patterns:
                match = re.search(pattern, line, re.IGNORECASE)
                if match:
                    date_str = match.group(1).strip()
                    # Convert date format if needed
                    result['invoice_date'] = normalize_date(date_str)
                    break
        
        # Extract supplier name (usually near the top)
        if not result['supplier_name'] and i < 15:
            supplier_patterns = [
                r'(?:supplier|furnizor|seller)[:\s]+(.+)',
                r'(?:s\.?c\.?)\s+([A-Z][A-Za-z\s&\.]+(?:s\.?r\.?l\.?|s\.?a\.?))',
            ]
            for pattern in supplier_patterns:
                match = re.search(pattern, line, re.IGNORECASE)
                if match:
                    result['supplier_name'] = match.group(1).strip()
                    break
            if not result['supplier_name'] and is_likely_supplier_line(line):
                result['supplier_name'] = line.strip()
        
        # Extract total amount
        if any(keyword in line_lower for keyword in TOTAL_KEYWORDS):
            for amount in extract_amount_candidates(line):
                if amount > result['total_amount']:
                    result['total_amount'] = amount
    
    # Fallback totals and currency detection from full text
    text_lower = text.lower()
    if "eur" in text_lower:
        result["currency"] = "EUR"
    elif "usd" in text_lower:
        result["currency"] = "USD"

    if result["total_amount"] <= 0:
        for amount in extract_amount_candidates(text):
            if amount > result["total_amount"]:
                result["total_amount"] = amount
    
    # Extract line items
    # Look for table-like structures with materials
    result['items'] = extract_line_items(text)
    
    return result


def extract_line_items(text: str) -> List[Dict]:
    """Extract line items from invoice text.
    
    This function looks for patterns that resemble invoice line items:
    - Description, quantity, unit price, total
    - Tab or space-separated values
    - Numeric patterns for quantities and prices
    """
    items = []
    lines = [line.strip() for line in text.split('\n')]
    
    # Try to identify table headers
    header_idx = -1
    for i, line in enumerate(lines):
        line_lower = line.lower()
        # Look for common table headers
        if any(keyword in line_lower for keyword in ['descriere', 'description', 'produs', 'material', 'item', 'denumire']):
            if any(keyword in line_lower for keyword in ['cantitate', 'quantity', 'qty', 'cant']):
                header_idx = i
                break
    
    # If we found a header, parse the following lines as table rows
    if header_idx >= 0:
        for i in range(header_idx + 1, min(header_idx + 100, len(lines))):
            line = lines[i]
            if not line:
                continue
            if any(stop in line.lower() for stop in ["subtotal", "total", "tva", "observatii", "notes"]):
                break
            
            # Try to extract item information using patterns
            item = parse_line_item(line)
            if item:
                items.append(item)
    
    # If no items found via table parsing, try a more aggressive approach
    if not items:
        for line in lines:
            item = parse_line_item(line)
            if item and item['quantity'] > 0 and item['unit_price'] > 0:
                items.append(item)

    # De-duplicate very similar rows (common in extracted PDFs)
    deduped: List[Dict] = []
    seen = set()
    for item in items:
        key = (
            item["description"].lower(),
            round(float(item["quantity"]), 4),
            round(float(item["unit_price"]), 4),
            round(float(item["total_price"]), 4),
        )
        if key in seen:
            continue
        seen.add(key)
        deduped.append(item)

    return deduped


def parse_line_item(line: str) -> Optional[Dict]:
    """Parse a single line item from invoice text.
    
    Expects patterns like:
    - "Product name    10    100.00    1000.00"
    - "Product name | 10 | 100.00 | 1000.00"
    - "Product name, 10 buc, 100.00 RON, 1000.00 RON"
    """
    line_lower = line.lower().strip()
    header_keywords = ['total', 'subtotal', 'tva', 'tax', 'discount', 'observatii', 'notes']
    table_header_keywords = ['descriere', 'description', 'cantitate', 'quantity', 'pret', 'price', 'produs', 'product', 'denumire']

    if any(keyword in line_lower for keyword in header_keywords):
        return None

    header_count = sum(1 for kw in table_header_keywords if kw in line_lower)
    if header_count >= 2 and len(re.findall(r'\d', line)) < 3:
        return None

    # Support both pipe/tabular and simple spaced rows.
    working_line = re.sub(r"[|;\t]", " ", line).strip()
    amount_matches = re.findall(
        r"(?<!\d)(?:\d{1,3}(?:[.\s,]\d{3})+|\d+)(?:[.,]\d{2})(?!\d)",
        working_line,
    )
    if len(amount_matches) < 1:
        return None

    amount_values = [v for v in (parse_decimal(x) for x in amount_matches) if v is not None]
    if not amount_values:
        return None

    # Determine unit/total from the right-most amounts.
    if len(amount_values) >= 2:
        unit_price = amount_values[-2]
        total_price = amount_values[-1]
    else:
        unit_price = amount_values[-1]
        total_price = amount_values[-1]

    # Quantity can be integer or decimal, optionally followed by unit labels.
    qty_match = re.search(
        r"(?<!\d)(\d+(?:[.,]\d{1,3})?)\s*(?:buc|kg|m|mp|set|pcs|x)?\b",
        working_line,
        re.IGNORECASE,
    )
    quantity = parse_decimal(qty_match.group(1)) if qty_match else None
    if quantity is None or quantity <= 0:
        # Fallback based on amount relation when possible.
        if unit_price > 0 and total_price >= unit_price:
            quantity = round(total_price / unit_price, 3)
        else:
            quantity = 1.0

    # Extract description as text before first numeric chunk.
    first_num_match = re.search(r'\d', working_line)
    if first_num_match:
        description = working_line[:first_num_match.start()].strip(" -.:")
    else:
        description = working_line[:80].strip()

    if len(description) < 3 or description.replace(' ', '').isdigit():
        return None
    if re.match(r"^(factura|invoice|data|furnizor|supplier)\b", description.lower()):
        return None

    return {
        'description': description,
        'sku': '',
        'quantity': float(quantity),
        'unit_price': float(unit_price),
        'total_price': float(total_price),
    }


def normalize_date(date_str: str) -> str:
    """Normalize date string to ISO format (YYYY-MM-DD)."""
    # Try different date formats
    patterns = [
        (r'(\d{1,2})[\./-](\d{1,2})[\./-](\d{4})', lambda m: f"{m.group(3)}-{m.group(2).zfill(2)}-{m.group(1).zfill(2)}"),
        (r'(\d{4})-(\d{2})-(\d{2})', lambda m: f"{m.group(1)}-{m.group(2)}-{m.group(3)}"),
    ]
    
    for pattern, formatter in patterns:
        match = re.match(pattern, date_str)
        if match:
            return formatter(match)
    
    return date_str


def extract_document_text(file_path: str, file_extension: str) -> str:
    """Extract raw text from supported document types."""
    # Extract text based on file type
    if file_extension == 'pdf':
        return extract_text_from_pdf(file_path)
    elif file_extension == 'docx':
        return extract_text_from_docx(file_path)
    elif file_extension == 'doc':
        return extract_text_from_doc(file_path)
    elif file_extension == 'txt':
        return extract_text_from_txt(file_path)
    raise Exception(f"Unsupported file format: {file_extension}")


def parse_document(file_path: str, file_extension: str) -> Dict:
    """Parse document and extract invoice data."""
    text = extract_document_text(file_path, file_extension)

    if not text:
        logger.warning(f"No text extracted from {file_extension} file")
        return EMPTY_PARSE_RESULT.copy()
    
    # Parse the extracted text
    return parse_invoice_materials(text)
